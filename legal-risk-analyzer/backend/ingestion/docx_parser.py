
from typing import Union, IO
import io
import re

try:
    from docx import Document
except Exception as exc:  # pragma: no cover - defensive import
    raise ImportError(
        "python-docx is required to use docx parsing utilities. Install with `pip install python-docx`."
    ) from exc


def extract_text_from_docx(file: Union[str, bytes, IO]) -> str:
    """Extracts paragraph text from a .docx file and returns a clean string.

    Args:
        file: One of:
            - ``str``: path to a .docx file
            - ``bytes``: raw .docx file bytes
            - file-like object opened in binary mode (e.g., ``io.BytesIO`` or an UploadFile.file)

    Returns:
        A cleaned string containing all non-empty paragraphs from the document,
        joined by a blank line (``\n\n``) between paragraphs. Leading/trailing
        whitespace is removed and consecutive blank lines are collapsed.

    Example:
        >>> text = extract_text_from_docx("/path/to/doc.docx")
    """

    # Load the document depending on input type
    if isinstance(file, str):
        doc = Document(file)
    elif isinstance(file, (bytes, bytearray)):
        doc = Document(io.BytesIO(file))
    else:
        # Assume file-like object
        # python-docx accepts file-like objects positioned at start
        try:
            # try to rewind if possible
            file.seek(0)
        except Exception:
            pass
        doc = Document(file)

    # Collect non-empty paragraphs and strip whitespace
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    # Join paragraphs with a blank line between them
    text = "\n\n".join(paragraphs)

    # Collapse 3+ newlines to two (preserve paragraph separation) and strip
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    return text


__all__ = ["extract_text_from_docx"]
